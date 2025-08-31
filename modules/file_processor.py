"""
File Processor - Advanced file and document processing system
"""

import asyncio
import logging
import json
import os
import mimetypes
from datetime import datetime
from typing import Dict, List, Any, Optional, Union, BinaryIO
from pathlib import Path
import hashlib
import zipfile
import tarfile
import tempfile

# Document processing imports
try:
    import PyPDF2
    import docx
    from openpyxl import load_workbook
    import csv
    from PIL import Image
    DOCUMENT_PROCESSING_AVAILABLE = True
except ImportError:
    DOCUMENT_PROCESSING_AVAILABLE = False

# OCR imports
try:
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)

class FileProcessor:
    """Advanced file and document processing system"""
    
    def __init__(self, llm_manager, settings):
        self.llm_manager = llm_manager
        self.settings = settings
        self.processing_history = []
        self.file_cache = {}
        self.extraction_rules = {}
        
        # Configuration
        self.max_file_size = 100 * 1024 * 1024  # 100MB
        self.temp_dir = Path("temp/file_processing")
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        
        # Supported file types
        self.supported_types = {
            'text': ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.css', '.js', '.py', '.log'],
            'document': ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.svg', '.webp'],
            'audio': ['.mp3', '.wav', '.ogg', '.m4a', '.flac', '.aac'],
            'video': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.webm', '.mkv'],
            'archive': ['.zip', '.rar', '.7z', '.tar', '.gz', '.bz2'],
            'code': ['.py', '.js', '.html', '.css', '.java', '.cpp', '.c', '.php', '.rb', '.go']
        }
        
    async def initialize(self):
        """Initialize file processor"""
        try:
            logger.info("Initializing File Processor...")
            
            # Check available libraries
            self.check_capabilities()
            
            # Load extraction rules
            await self.load_extraction_rules()
            
            # Setup file watchers
            await self.setup_file_watchers()
            
            logger.info("File Processor initialized")
            
        except Exception as e:
            logger.error(f"Failed to initialize File Processor: {e}")
            raise
    
    def check_capabilities(self):
        """Check what file processing capabilities are available"""
        self.capabilities = {
            'document_processing': DOCUMENT_PROCESSING_AVAILABLE,
            'ocr': OCR_AVAILABLE,
            'image_processing': True,  # PIL is usually available
            'text_analysis': True,
            'metadata_extraction': True
        }
        
        logger.info(f"File processing capabilities: {self.capabilities}")
    
    async def process_action(self, action_params: Dict[str, Any]) -> Dict[str, Any]:
        """Process file-related actions"""
        try:
            action_type = action_params.get('action_type')
            result = {'action': action_type, 'success': False, 'timestamp': datetime.now().isoformat()}
            
            if action_type == 'analyze_file':
                result = await self.analyze_file(action_params)
            elif action_type == 'extract_text':
                result = await self.extract_text(action_params)
            elif action_type == 'convert_file':
                result = await self.convert_file(action_params)
            elif action_type == 'batch_process':
                result = await self.batch_process(action_params)
            elif action_type == 'organize_files':
                result = await self.organize_files(action_params)
            elif action_type == 'search_content':
                result = await self.search_content(action_params)
            elif action_type == 'compress_files':
                result = await self.compress_files(action_params)
            elif action_type == 'extract_archive':
                result = await self.extract_archive(action_params)
            elif action_type == 'generate_report':
                result = await self.generate_file_report(action_params)
            else:
                result['error'] = f'Unknown action type: {action_type}'
            
            # Store in history
            self.processing_history.append(result)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing file action: {e}")
            return {'action': action_type, 'success': False, 'error': str(e)}
    
    async def analyze_file(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze a file and extract comprehensive information"""
        try:
            file_path = params.get('file_path', '')
            if not file_path or not os.path.exists(file_path):
                return {'success': False, 'error': 'File not found'}
            
            file_path = Path(file_path)
            
            # Basic file information
            file_info = await self.get_file_info(file_path)
            
            # Content analysis based on file type
            content_analysis = await self.analyze_file_content(file_path)
            
            # Metadata extraction
            metadata = await self.extract_metadata(file_path)
            
            # Security analysis
            security_info = await self.analyze_file_security(file_path)
            
            # AI-powered analysis if text content available
            ai_analysis = None
            if content_analysis.get('text_content'):
                ai_analysis = await self.ai_analyze_content(content_analysis['text_content'])
            
            return {
                'success': True,
                'file_info': file_info,
                'content_analysis': content_analysis,
                'metadata': metadata,
                'security_info': security_info,
                'ai_analysis': ai_analysis
            }
            
        except Exception as e:
            logger.error(f"Error analyzing file: {e}")
            return {'success': False, 'error': str(e)}
    
    async def get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get basic file information"""
        try:
            stat = file_path.stat()
            
            return {
                'name': file_path.name,
                'path': str(file_path),
                'size': stat.st_size,
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'extension': file_path.suffix.lower(),
                'mime_type': mimetypes.guess_type(str(file_path))[0],
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed': datetime.fromtimestamp(stat.st_atime).isoformat(),
                'is_readonly': not os.access(file_path, os.W_OK),
                'file_type': self.categorize_file_type(file_path.suffix.lower())
            }
            
        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return {'error': str(e)}
    
    def categorize_file_type(self, extension: str) -> str:
        """Categorize file type based on extension"""
        for category, extensions in self.supported_types.items():
            if extension in extensions:
                return category
        return 'unknown'
    
    async def analyze_file_content(self, file_path: Path) -> Dict[str, Any]:
        """Analyze file content based on type"""
        try:
            file_type = self.categorize_file_type(file_path.suffix.lower())
            
            if file_type == 'text':
                return await self.analyze_text_file(file_path)
            elif file_type == 'document':
                return await self.analyze_document_file(file_path)
            elif file_type == 'image':
                return await self.analyze_image_file(file_path)
            elif file_type == 'code':
                return await self.analyze_code_file(file_path)
            elif file_type == 'archive':
                return await self.analyze_archive_file(file_path)
            else:
                return {'type': file_type, 'analysis': 'Content analysis not available for this file type'}
                
        except Exception as e:
            logger.error(f"Error analyzing file content: {e}")
            return {'error': str(e)}
    
    async def analyze_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze text-based files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Basic text statistics
            lines = content.split('\n')
            words = content.split()
            characters = len(content)
            
            # Language detection (simple heuristic)
            language = self.detect_language(content)
            
            # Extract keywords
            keywords = self.extract_keywords(content)
            
            return {
                'type': 'text',
                'encoding': 'utf-8',
                'line_count': len(lines),
                'word_count': len(words),
                'character_count': characters,
                'language': language,
                'keywords': keywords,
                'text_content': content[:5000],  # First 5000 chars
                'is_empty': len(content.strip()) == 0,
                'has_code': self.detect_code_patterns(content)
            }
            
        except Exception as e:
            logger.error(f"Error analyzing text file: {e}")
            return {'error': str(e)}
    
    async def analyze_document_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze document files (PDF, Word, Excel, etc.)"""
        try:
            if not DOCUMENT_PROCESSING_AVAILABLE:
                return {'error': 'Document processing libraries not available'}
            
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return await self.analyze_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return await self.analyze_word_document(file_path)
            elif extension in ['.xlsx', '.xls']:
                return await self.analyze_excel_document(file_path)
            else:
                return {'type': 'document', 'analysis': f'Analysis not implemented for {extension}'}
                
        except Exception as e:
            logger.error(f"Error analyzing document file: {e}")
            return {'error': str(e)}
    
    async def analyze_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Analyze PDF files"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                page_count = len(pdf_reader.pages)
                text_content = ""
                
                # Extract text from all pages
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                
                # Get metadata
                metadata = pdf_reader.metadata if pdf_reader.metadata else {}
                
                return {
                    'type': 'pdf',
                    'page_count': page_count,
                    'text_content': text_content[:5000],  # First 5000 chars
                    'word_count': len(text_content.split()),
                    'has_text': len(text_content.strip()) > 0,
                    'metadata': {
                        'title': metadata.get('/Title', ''),
                        'author': metadata.get('/Author', ''),
                        'subject': metadata.get('/Subject', ''),
                        'creator': metadata.get('/Creator', ''),
                        'producer': metadata.get('/Producer', ''),
                        'creation_date': str(metadata.get('/CreationDate', '')),
                        'modification_date': str(metadata.get('/ModDate', ''))
                    }
                }
                
        except Exception as e:
            logger.error(f"Error analyzing PDF: {e}")
            return {'error': str(e)}
    
    async def analyze_word_document(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Word documents"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Count elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Get core properties
            core_props = doc.core_properties
            
            return {
                'type': 'word_document',
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'text_content': text_content[:5000],
                'word_count': len(text_content.split()),
                'metadata': {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else '',
                    'last_modified_by': core_props.last_modified_by or '',
                    'revision': core_props.revision or ''
                }
            }
            
        except Exception as e:
            logger.error(f"Error analyzing Word document: {e}")
            return {'error': str(e)}
    
    async def analyze_excel_document(self, file_path: Path) -> Dict[str, Any]:
        """Analyze Excel spreadsheets"""
        try:
            workbook = load_workbook(file_path, data_only=True)
            
            sheet_info = []
            total_rows = 0
            total_cols = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                sheet_info.append({
                    'name': sheet_name,
                    'rows': max_row,
                    'columns': max_col,
                    'cells': max_row * max_col
                })
                
                total_rows += max_row
                total_cols = max(total_cols, max_col)
            
            # Extract some sample data
            sample_data = []
            if workbook.active:
                for row in workbook.active.iter_rows(min_row=1, max_row=5, values_only=True):
                    sample_data.append([str(cell) if cell is not None else '' for cell in row])
            
            return {
                'type': 'excel_spreadsheet',
                'sheet_count': len(workbook.sheetnames),
                'sheet_info': sheet_info,
                'total_rows': total_rows,
                'max_columns': total_cols,
                'sample_data': sample_data,
                'metadata': {
                    'properties': {
                        'title': workbook.properties.title or '',
                        'author': workbook.properties.creator or '',
                        'subject': workbook.properties.subject or '',
                        'created': workbook.properties.created.isoformat() if workbook.properties.created else '',
                        'modified': workbook.properties.modified.isoformat() if workbook.properties.modified else ''
                    }
                }
            }
            
        except Exception as e:
            logger.error(f"Error analyzing Excel document: {e}")
            return {'error': str(e)}
    
    async def analyze_image_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze image files"""
        try:
            with Image.open(file_path) as img:
                # Basic image info
                width, height = img.size
                mode = img.mode
                format_type = img.format
                
                # Calculate file statistics
                aspect_ratio = width / height
                megapixels = (width * height) / 1000000
                
                # OCR if available
                ocr_text = ""
                if OCR_AVAILABLE:
                    try:
                        ocr_text = pytesseract.image_to_string(img)
                    except Exception as e:
                        logger.warning(f"OCR failed: {e}")
                
                # Color analysis
                colors = self.analyze_image_colors(img)
                
                return {
                    'type': 'image',
                    'width': width,
                    'height': height,
                    'mode': mode,
                    'format': format_type,
                    'aspect_ratio': round(aspect_ratio, 2),
                    'megapixels': round(megapixels, 2),
                    'has_transparency': mode in ['RGBA', 'LA'] or 'transparency' in img.info,
                    'ocr_text': ocr_text[:1000] if ocr_text else '',
                    'dominant_colors': colors,
                    'estimated_quality': self.estimate_image_quality(img)
                }
                
        except Exception as e:
            logger.error(f"Error analyzing image file: {e}")
            return {'error': str(e)}
    
    def analyze_image_colors(self, img: Image.Image) -> List[str]:
        """Analyze dominant colors in an image"""
        try:
            # Convert to RGB if necessary
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize for faster processing
            img_small = img.resize((50, 50))
            
            # Get pixel data and find most common colors
            pixels = list(img_small.getdata())
            color_counts = {}
            
            for pixel in pixels:
                color = f"#{pixel[0]:02x}{pixel[1]:02x}{pixel[2]:02x}"
                color_counts[color] = color_counts.get(color, 0) + 1
            
            # Return top 5 colors
            sorted_colors = sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
            return [color for color, count in sorted_colors[:5]]
            
        except Exception as e:
            logger.error(f"Error analyzing image colors: {e}")
            return []
    
    def estimate_image_quality(self, img: Image.Image) -> str:
        """Estimate image quality based on various factors"""
        try:
            width, height = img.size
            total_pixels = width * height
            
            if total_pixels > 8000000:  # > 8MP
                return "high"
            elif total_pixels > 2000000:  # > 2MP
                return "medium"
            else:
                return "low"
                
        except Exception as e:
            logger.error(f"Error estimating image quality: {e}")
            return "unknown"
    
    async def analyze_code_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze code files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            # Basic code statistics
            total_lines = len(lines)
            code_lines = len([line for line in lines if line.strip() and not line.strip().startswith('#')])
            comment_lines = len([line for line in lines if line.strip().startswith('#')])
            blank_lines = total_lines - code_lines - comment_lines
            
            # Language detection
            language = self.detect_programming_language(file_path.suffix, content)
            
            # Code complexity analysis
            complexity = self.analyze_code_complexity(content, language)
            
            return {
                'type': 'code',
                'language': language,
                'total_lines': total_lines,
                'code_lines': code_lines,
                'comment_lines': comment_lines,
                'blank_lines': blank_lines,
                'comment_ratio': round(comment_lines / max(total_lines, 1), 2),
                'complexity': complexity,
                'text_content': content[:5000],  # First 5000 chars
                'functions': self.extract_function_names(content, language),
                'imports': self.extract_imports(content, language)
            }
            
        except Exception as e:
            logger.error(f"Error analyzing code file: {e}")
            return {'error': str(e)}
    
    def detect_programming_language(self, extension: str, content: str) -> str:
        """Detect programming language from file extension and content"""
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.html': 'html',
            '.css': 'css',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.sh': 'bash'
        }
        
        return language_map.get(extension, 'unknown')
    
    def analyze_code_complexity(self, content: str, language: str) -> Dict[str, Any]:
        """Analyze code complexity"""
        try:
            lines = content.split('\n')
            
            # Count control structures
            control_keywords = ['if', 'for', 'while', 'switch', 'case', 'try', 'catch']
            complexity_score = 0
            
            for line in lines:
                line_lower = line.lower().strip()
                for keyword in control_keywords:
                    if keyword in line_lower:
                        complexity_score += 1
            
            # Determine complexity level
            if complexity_score < 10:
                level = "low"
            elif complexity_score < 30:
                level = "medium"
            else:
                level = "high"
            
            return {
                'score': complexity_score,
                'level': level,
                'control_structures': complexity_score
            }
            
        except Exception as e:
            logger.error(f"Error analyzing code complexity: {e}")
            return {'score': 0, 'level': 'unknown'}
    
    def extract_function_names(self, content: str, language: str) -> List[str]:
        """Extract function names from code"""
        try:
            import re
            functions = []
            
            if language == 'python':
                pattern = r'def\s+(\w+)\s*\('
                functions = re.findall(pattern, content)
            elif language == 'javascript':
                pattern = r'function\s+(\w+)\s*\('
                functions = re.findall(pattern, content)
            elif language == 'java':
                pattern = r'(?:public|private|protected)?\s*\w+\s+(\w+)\s*\('
                functions = re.findall(pattern, content)
            
            return functions[:20]  # Return first 20 functions
            
        except Exception as e:
            logger.error(f"Error extracting function names: {e}")
            return []
    
    def extract_imports(self, content: str, language: str) -> List[str]:
        """Extract import statements from code"""
        try:
            import re
            imports = []
            
            if language == 'python':
                pattern = r'(?:from\s+\S+\s+)?import\s+([^\n]+)'
                imports = re.findall(pattern, content)
            elif language == 'javascript':
                pattern = r'import\s+.*?from\s+[\'"]([^\'"]+)[\'"]'
                imports = re.findall(pattern, content)
            elif language == 'java':
                pattern = r'import\s+([^\s;]+);'
                imports = re.findall(pattern, content)
            
            return imports[:10]  # Return first 10 imports
            
        except Exception as e:
            logger.error(f"Error extracting imports: {e}")
            return []
    
    async def analyze_archive_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze archive files"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.zip':
                return await self.analyze_zip_file(file_path)
            elif extension in ['.tar', '.gz', '.bz2']:
                return await self.analyze_tar_file(file_path)
            else:
                return {'type': 'archive', 'analysis': f'Archive type {extension} not supported'}
                
        except Exception as e:
            logger.error(f"Error analyzing archive file: {e}")
            return {'error': str(e)}
    
    async def analyze_zip_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze ZIP files"""
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                file_list = zip_file.namelist()
                file_count = len(file_list)
                
                # Analyze file types in archive
                file_types = {}
                total_uncompressed_size = 0
                
                for file_info in zip_file.infolist():
                    if not file_info.is_dir():
                        ext = Path(file_info.filename).suffix.lower()
                        file_types[ext] = file_types.get(ext, 0) + 1
                        total_uncompressed_size += file_info.file_size
                
                compression_ratio = 1 - (file_path.stat().st_size / max(total_uncompressed_size, 1))
                
                return {
                    'type': 'zip_archive',
                    'file_count': file_count,
                    'file_list': file_list[:50],  # First 50 files
                    'file_types': file_types,
                    'uncompressed_size': total_uncompressed_size,
                    'compression_ratio': round(compression_ratio, 2),
                    'has_subdirectories': any('/' in name for name in file_list)
                }
                
        except Exception as e:
            logger.error(f"Error analyzing ZIP file: {e}")
            return {'error': str(e)}
    
    async def analyze_tar_file(self, file_path: Path) -> Dict[str, Any]:
        """Analyze TAR files"""
        try:
            with tarfile.open(file_path, 'r') as tar_file:
                members = tar_file.getmembers()
                file_count = len([m for m in members if m.isfile()])
                dir_count = len([m for m in members if m.isdir()])
                
                # Analyze file types
                file_types = {}
                total_size = 0
                
                for member in members:
                    if member.isfile():
                        ext = Path(member.name).suffix.lower()
                        file_types[ext] = file_types.get(ext, 0) + 1
                        total_size += member.size
                
                return {
                    'type': 'tar_archive',
                    'file_count': file_count,
                    'directory_count': dir_count,
                    'file_types': file_types,
                    'total_size': total_size,
                    'members': [m.name for m in members[:50]]  # First 50 members
                }
                
        except Exception as e:
            logger.error(f"Error analyzing TAR file: {e}")
            return {'error': str(e)}
    
    async def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract file metadata"""
        try:
            metadata = {
                'file_hash': await self.calculate_file_hash(file_path),
                'checksum': await self.calculate_checksum(file_path)
            }
            
            # Extended attributes (if supported)
            try:
                import xattr
                extended_attrs = dict(xattr.xattr(file_path))
                metadata['extended_attributes'] = extended_attrs
            except ImportError:
                pass
            except Exception as e:
                logger.debug(f"Could not read extended attributes: {e}")
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")
            return {'error': str(e)}
    
    async def calculate_file_hash(self, file_path: Path, algorithm: str = 'sha256') -> str:
        """Calculate file hash"""
        try:
            hash_obj = hashlib.new(algorithm)
            
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_obj.update(chunk)
            
            return hash_obj.hexdigest()
            
        except Exception as e:
            logger.error(f"Error calculating file hash: {e}")
            return ""
    
    async def calculate_checksum(self, file_path: Path) -> str:
        """Calculate file checksum (MD5)"""
        try:
            md5_hash = hashlib.md5()
            
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    md5_hash.update(chunk)
            
            return md5_hash.hexdigest()
            
        except Exception as e:
            logger.error(f"Error calculating checksum: {e}")
            return ""
    
    async def analyze_file_security(self, file_path: Path) -> Dict[str, Any]:
        """Analyze file for security issues"""
        try:
            security_info = {
                'potentially_dangerous': False,
                'reasons': [],
                'file_type_risk': 'low'
            }
            
            extension = file_path.suffix.lower()
            
            # Check for potentially dangerous file types
            dangerous_extensions = ['.exe', '.bat', '.cmd', '.scr', '.pif', '.com', '.jar']
            if extension in dangerous_extensions:
                security_info['potentially_dangerous'] = True
                security_info['reasons'].append(f'Executable file type: {extension}')
                security_info['file_type_risk'] = 'high'
            
            # Check file size (unusually large files might be suspicious)
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size:
                security_info['reasons'].append(f'Large file size: {file_size} bytes')
            
            # Check for hidden file attributes
            if file_path.name.startswith('.'):
                security_info['reasons'].append('Hidden file')
            
            return security_info
            
        except Exception as e:
            logger.error(f"Error analyzing file security: {e}")
            return {'error': str(e)}
    
    async def ai_analyze_content(self, text_content: str) -> Dict[str, Any]:
        """Use AI to analyze text content"""
        try:
            if not text_content.strip():
                return {'analysis': 'No text content to analyze'}
            
            # Create analysis prompt
            prompt = f"""
            Analyze the following text content and provide insights:
            
            Content (first 2000 characters):
            {text_content[:2000]}
            
            Please provide:
            1. Content summary
            2. Main topics or themes
            3. Language and tone analysis
            4. Key information extracted
            5. Content classification (document type, purpose)
            
            Respond in JSON format with structured analysis.
            """
            
            # Get AI analysis
            ai_response = await self.llm_manager.generate_response(prompt)
            
            try:
                # Try to parse as JSON
                analysis = json.loads(ai_response)
                return analysis
            except json.JSONDecodeError:
                # If not valid JSON, return as text
                return {'analysis': ai_response}
            
        except Exception as e:
            logger.error(f"Error in AI content analysis: {e}")
            return {'error': str(e)}
    
    def detect_language(self, text: str) -> str:
        """Detect language of text content"""
        try:
            # Simple language detection based on common words
            english_words = ['the', 'and', 'is', 'in', 'to', 'of', 'a', 'that', 'it', 'with']
            spanish_words = ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'es', 'se', 'no']
            french_words = ['le', 'de', 'et', 'à', 'un', 'il', 'être', 'et', 'en', 'avoir']
            
            text_lower = text.lower()
            
            english_count = sum(1 for word in english_words if word in text_lower)
            spanish_count = sum(1 for word in spanish_words if word in text_lower)
            french_count = sum(1 for word in french_words if word in text_lower)
            
            if english_count >= spanish_count and english_count >= french_count:
                return 'english'
            elif spanish_count >= french_count:
                return 'spanish'
            elif french_count > 0:
                return 'french'
            else:
                return 'unknown'
                
        except Exception as e:
            logger.error(f"Error detecting language: {e}")
            return 'unknown'
    
    def extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from text"""
        try:
            # Simple keyword extraction
            import re
            
            # Remove common stop words
            stop_words = {
                'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
                'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before',
                'after', 'above', 'below', 'between', 'among', 'since', 'without',
                'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has',
                'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should',
                'a', 'an', 'this', 'that', 'these', 'those', 'i', 'you', 'he',
                'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them'
            }
            
            # Extract words
            words = re.findall(r'\b\w+\b', text.lower())
            
            # Filter out stop words and short words
            keywords = [word for word in words if len(word) > 3 and word not in stop_words]
            
            # Count word frequency
            word_count = {}
            for word in keywords:
                word_count[word] = word_count.get(word, 0) + 1
            
            # Sort by frequency and return top keywords
            sorted_keywords = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
            
            return [word for word, count in sorted_keywords[:max_keywords]]
            
        except Exception as e:
            logger.error(f"Error extracting keywords: {e}")
            return []
    
    def detect_code_patterns(self, text: str) -> bool:
        """Detect if text contains code patterns"""
        try:
            code_indicators = [
                'def ', 'function ', 'class ', 'import ', 'include ',
                '<?php', '#!/', 'public class', 'private ', 'public ',
                '{', '}', '()', '=>', '==', '!=', '&&', '||'
            ]
            
            text_lower = text.lower()
            return any(indicator in text_lower for indicator in code_indicators)
            
        except Exception as e:
            logger.error(f"Error detecting code patterns: {e}")
            return False
    
    async def extract_text(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from various file formats"""
        try:
            file_path = params.get('file_path', '')
            if not file_path or not os.path.exists(file_path):
                return {'success': False, 'error': 'File not found'}
            
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            
            extracted_text = ""
            
            if extension == '.pdf':
                extracted_text = await self.extract_text_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                extracted_text = await self.extract_text_from_word(file_path)
            elif extension in ['.txt', '.md', '.csv']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
            elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                if OCR_AVAILABLE:
                    extracted_text = await self.extract_text_from_image(file_path)
                else:
                    return {'success': False, 'error': 'OCR not available'}
            else:
                return {'success': False, 'error': f'Text extraction not supported for {extension}'}
            
            return {
                'success': True,
                'extracted_text': extracted_text,
                'text_length': len(extracted_text),
                'word_count': len(extracted_text.split()),
                'file_path': str(file_path)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return {'success': False, 'error': str(e)}
    
    async def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            if not DOCUMENT_PROCESSING_AVAILABLE:
                raise Exception("PDF processing not available")
            
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""
    
    async def extract_text_from_word(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            if not DOCUMENT_PROCESSING_AVAILABLE:
                raise Exception("Word processing not available")
            
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from Word: {e}")
            return ""
    
    async def extract_text_from_image(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        try:
            if not OCR_AVAILABLE:
                raise Exception("OCR not available")
            
            with Image.open(file_path) as img:
                text = pytesseract.image_to_string(img)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from image: {e}")
            return ""
    
    async def load_extraction_rules(self):
        """Load file extraction rules"""
        try:
            rules_file = Path("data/extraction_rules.json")
            if rules_file.exists():
                with open(rules_file, 'r') as f:
                    self.extraction_rules = json.load(f)
            else:
                # Default extraction rules
                self.extraction_rules = {
                    'auto_extract': True,
                    'max_file_size': self.max_file_size,
                    'supported_types': list(self.supported_types.keys())
                }
            
            logger.info(f"Loaded extraction rules with {len(self.extraction_rules)} settings")
            
        except Exception as e:
            logger.error(f"Error loading extraction rules: {e}")
    
    async def setup_file_watchers(self):
        """Setup file system watchers for automatic processing"""
        try:
            # File watching would be implemented here using libraries like watchdog
            logger.info("File watchers setup (placeholder)")
        except Exception as e:
            logger.error(f"Error setting up file watchers: {e}")
    
    async def cleanup_temp_files(self):
        """Clean up temporary files"""
        try:
            import shutil
            
            if self.temp_dir.exists():
                shutil.rmtree(self.temp_dir)
                self.temp_dir.mkdir(parents=True, exist_ok=True)
                
            logger.info("Temporary files cleaned up")
            
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")
    
    def get_processing_history(self, limit: int = 50) -> List[Dict[str, Any]]:
        """Get file processing history"""
        return self.processing_history[-limit:] if self.processing_history else []
    
    def get_supported_file_types(self) -> Dict[str, List[str]]:
        """Get supported file types"""
        return self.supported_types
    
    def get_capabilities_info(self) -> Dict[str, Any]:
        """Get information about processing capabilities"""
        return {
            'capabilities': self.capabilities,
            'supported_types': self.supported_types,
            'max_file_size_mb': self.max_file_size / (1024 * 1024),
            'temp_directory': str(self.temp_dir)
        }
    
    async def shutdown(self):
        """Shutdown file processor"""
        logger.info("Shutting down File Processor...")
        
        # Clean up temporary files
        await self.cleanup_temp_files()
        
        # Save processing history
        try:
            history_file = Path("data/file_processing_history.json")
            history_file.parent.mkdir(exist_ok=True)
            
            with open(history_file, 'w') as f:
                json.dump(self.processing_history[-100:], f, indent=2)  # Save last 100 items
                
        except Exception as e:
            logger.error(f"Error saving processing history: {e}")
